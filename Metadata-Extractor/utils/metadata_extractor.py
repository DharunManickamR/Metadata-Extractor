import os
import mimetypes
import re
from datetime import datetime
from PIL import Image
from PIL.ExifTags import TAGS, GPSTAGS
from pypdf import PdfReader
import docx
from tinytag import TinyTag
from hachoir.parser import createParser
from hachoir.metadata import extractMetadata

# Initialize mimetypes database
mimetypes.init()

# Define common categories
CATEGORY_MAP = {
    'image/jpeg': 'image',
    'image/png': 'image',
    'image/gif': 'image',
    'image/bmp': 'image',
    'image/webp': 'image',
    'application/pdf': 'document',
    'application/vnd.openxmlformats-officedocument.wordprocessingml.document': 'document',
    'audio/mpeg': 'audio',
    'audio/wav': 'audio',
    'audio/x-wav': 'audio',
    'audio/ogg': 'audio',
    'audio/flac': 'audio',
    'audio/mp4': 'audio',
    'video/mp4': 'video',
    'video/x-msvideo': 'video', # AVI
    'video/x-matroska': 'video', # MKV
    'video/quicktime': 'video',  # MOV
    'video/x-ms-wmv': 'video'    # WMV
}

def format_size(bytes_size):
    """Formats bytes size into human-readable format."""
    try:
        bytes_size = float(bytes_size)
    except (TypeError, ValueError):
        return "Unknown"
        
    for unit in ['Bytes', 'KB', 'MB', 'GB']:
        if bytes_size < 1024.0:
            return f"{bytes_size:.2f} {unit}"
        bytes_size /= 1024.0
    return f"{bytes_size:.2f} TB"

def get_file_category(mime_type, filename):
    """Categorizes the file based on MIME type and file extension."""
    if mime_type in CATEGORY_MAP:
        return CATEGORY_MAP[mime_type]
        
    ext = os.path.splitext(filename)[1].lower()
    if ext in ['.jpg', '.jpeg', '.png', '.gif', '.bmp', '.webp']:
        return 'image'
    elif ext in ['.pdf', '.docx']:
        return 'document'
    elif ext in ['.mp3', '.wav', '.ogg', '.flac', '.m4a']:
        return 'audio'
    elif ext in ['.mp4', '.avi', '.mkv', '.mov', '.wmv', '.flv']:
        return 'video'
    return 'unknown'

def parse_pdf_date(date_str):
    """Parses standard PDF date strings like D:20231024153022-05'00' or 20231024153022Z."""
    if not date_str:
        return None
    if date_str.startswith('D:'):
        date_str = date_str[2:]
        
    # Strip non-numeric/timezone chars
    date_str = re.sub(r"[^0-9TZ+\-']", "", date_str)
    
    # Try YYYYMMDDHHmmSS
    match = re.match(r"^(\d{4})(\d{2})(\d{2})(\d{2})(\d{2})(\d{2})", date_str)
    if match:
        parts = match.groups()
        try:
            dt = datetime(int(parts[0]), int(parts[1]), int(parts[2]), int(parts[3]), int(parts[4]), int(parts[5]))
            return dt.strftime("%Y-%m-%d %H:%M:%S")
        except ValueError:
            pass
            
    # Try YYYYMMDD
    match_short = re.match(r"^(\d{4})(\d{2})(\d{2})", date_str)
    if match_short:
        parts = match_short.groups()
        try:
            dt = datetime(int(parts[0]), int(parts[1]), int(parts[2]))
            return dt.strftime("%Y-%m-%d")
        except ValueError:
            pass
            
    return date_str

def parse_gps_coordinates(gps_info):
    """Converts DMS GPS coordinates from PIL EXIF to decimal degrees."""
    if not gps_info:
        return None
        
    def to_decimal(rational_coords, ref):
        if not rational_coords or not ref:
            return None
        try:
            # Handle list/tuple of floats, ints, or rational numbers
            d = float(rational_coords[0])
            m = float(rational_coords[1])
            s = float(rational_coords[2])
            dec = d + (m / 60.0) + (s / 3600.0)
            if ref in ['S', 'W']:
                dec = -dec
            return dec
        except Exception:
            return None

    lat_ref = gps_info.get("GPSLatitudeRef")
    lat = gps_info.get("GPSLatitude")
    lon_ref = gps_info.get("GPSLongitudeRef")
    lon = gps_info.get("GPSLongitude")
    
    latitude = to_decimal(lat, lat_ref)
    longitude = to_decimal(lon, lon_ref)
    
    altitude_val = gps_info.get("GPSAltitude")
    altitude_ref = gps_info.get("GPSAltitudeRef")
    altitude = None
    if altitude_val is not None:
        try:
            altitude = float(altitude_val)
            if altitude_ref == 1 or (isinstance(altitude_ref, bytes) and altitude_ref == b'\x01'):
                altitude = -altitude
        except Exception:
            pass
            
    if latitude is not None and longitude is not None:
        return {
            "latitude": latitude,
            "longitude": longitude,
            "altitude": altitude,
            "maps_url": f"https://www.google.com/maps/search/?api=1&query={latitude},{longitude}"
        }
    return None

def extract_image_metadata(filepath):
    """Extracts resolution, format, EXIF and GPS from images."""
    specific = {}
    raw = {}
    try:
        with Image.open(filepath) as img:
            specific["Resolution"] = f"{img.width} x {img.height}"
            specific["Format"] = img.format
            specific["Color Mode"] = img.mode
            
            raw["resolution_width"] = img.width
            raw["resolution_height"] = img.height
            raw["format"] = img.format
            
            # Extract EXIF tags
            exif = img.getexif()
            if exif:
                for tag, value in exif.items():
                    tag_name = TAGS.get(tag, tag)
                    if tag_name == "GPSInfo":
                        gps_data = {}
                        for key in value:
                            sub_tag = GPSTAGS.get(key, key)
                            gps_data[sub_tag] = value[key]
                        raw["gps_info_raw"] = {str(k): str(v) for k, v in gps_data.items()}
                        gps_coords = parse_gps_coordinates(gps_data)
                        if gps_coords:
                            specific["GPS Latitude"] = f"{gps_coords['latitude']:.6f}"
                            specific["GPS Longitude"] = f"{gps_coords['longitude']:.6f}"
                            if gps_coords['altitude'] is not None:
                                specific["GPS Altitude"] = f"{gps_coords['altitude']:.2f} m"
                            specific["Google Maps Link"] = gps_coords["maps_url"]
                    else:
                        if isinstance(value, bytes):
                            try:
                                value = value.decode('utf-8', errors='ignore').strip()
                            except Exception:
                                pass
                        raw[tag_name] = str(value)
                        
            # Extract secondary EXIF info (more camera specific details)
            if hasattr(img, '_getexif'):
                _exif = img._getexif()
                if _exif:
                    for tag, value in _exif.items():
                        tag_name = TAGS.get(tag, tag)
                        if tag_name != "GPSInfo":
                            if isinstance(value, bytes):
                                try:
                                    value = value.decode('utf-8', errors='ignore').strip()
                                except Exception:
                                    pass
                            raw[tag_name] = str(value)
                            
            # Promote specific key camera tags to specific fields
            camera_mappings = {
                "Make": "Camera Manufacturer",
                "Model": "Camera Model",
                "Software": "Software Used",
                "DateTimeOriginal": "Capture Date/Time",
                "ExposureTime": "Exposure Time",
                "FNumber": "F-Number",
                "ISOSpeedRatings": "ISO Speed",
                "ShutterSpeedValue": "Shutter Speed",
                "FocalLength": "Focal Length"
            }
            for raw_key, label in camera_mappings.items():
                if raw_key in raw:
                    specific[label] = raw[raw_key]
                    
    except Exception as e:
        specific["Error"] = f"Failed to extract image metadata: {str(e)}"
        
    return specific, raw

def extract_pdf_metadata(filepath):
    """Extracts page count, author, dates and titles from PDF files."""
    specific = {}
    raw = {}
    try:
        reader = PdfReader(filepath)
        num_pages = len(reader.pages)
        specific["Pages"] = str(num_pages)
        raw["pages"] = num_pages
        
        info = reader.metadata
        if info:
            mappings = {
                "author": ("Author", "/Author"),
                "creator": ("Creator Software", "/Creator"),
                "producer": ("PDF Producer", "/Producer"),
                "subject": ("Subject", "/Subject"),
                "title": ("Title", "/Title"),
                "keywords": ("Keywords", "/Keywords")
            }
            
            for key, (label, pdf_key) in mappings.items():
                val = info.get(pdf_key) or getattr(info, key, None)
                if val:
                    specific[label] = str(val)
                    raw[key] = str(val)
                    
            creation_date = info.get('/CreationDate')
            mod_date = info.get('/ModDate')
            
            if creation_date:
                raw["creation_date_raw"] = str(creation_date)
                parsed_c = parse_pdf_date(str(creation_date))
                if parsed_c:
                    specific["Creation Date"] = parsed_c
                    
            if mod_date:
                raw["modification_date_raw"] = str(mod_date)
                parsed_m = parse_pdf_date(str(mod_date))
                if parsed_m:
                    specific["Modification Date"] = parsed_m
                    
    except Exception as e:
        specific["Error"] = f"Failed to extract PDF metadata: {str(e)}"
        
    return specific, raw

def extract_docx_metadata(filepath):
    """Extracts Word document core properties."""
    specific = {}
    raw = {}
    try:
        doc = docx.Document(filepath)
        prop = doc.core_properties
        
        mappings = {
            "title": "Title",
            "subject": "Subject",
            "author": "Author",
            "keywords": "Keywords",
            "comments": "Comments",
            "last_modified_by": "Last Modified By",
            "revision": "Revision"
        }
        
        for attr, label in mappings.items():
            val = getattr(prop, attr, None)
            if val:
                specific[label] = str(val)
                raw[attr] = str(val)
                
        if prop.created:
            specific["Creation Date"] = prop.created.strftime("%Y-%m-%d %H:%M:%S")
            raw["created"] = prop.created.isoformat()
            
        if prop.modified:
            specific["Modification Date"] = prop.modified.strftime("%Y-%m-%d %H:%M:%S")
            raw["modified"] = prop.modified.isoformat()
            
        # Add basic count information
        specific["Paragraphs Count"] = str(len(doc.paragraphs))
        specific["Tables Count"] = str(len(doc.tables))
        raw["paragraphs_count"] = len(doc.paragraphs)
        raw["tables_count"] = len(doc.tables)
        
    except Exception as e:
        specific["Error"] = f"Failed to extract DOCX metadata: {str(e)}"
        
    return specific, raw

def extract_audio_video_metadata(filepath):
    """Extracts duration, bitrate, sample rates, and codecs using tinytag & hachoir."""
    specific = {}
    raw = {}
    
    # 1. Try tinytag for audio/music metadata first
    try:
        tag = TinyTag.get(filepath)
        if tag:
            if tag.title:
                specific["Title"] = tag.title
                raw["title"] = tag.title
            if tag.artist:
                specific["Artist"] = tag.artist
                raw["artist"] = tag.artist
            if tag.album:
                specific["Album"] = tag.album
                raw["album"] = tag.album
            if tag.genre:
                specific["Genre"] = tag.genre
                raw["genre"] = tag.genre
            if tag.year:
                specific["Year"] = str(tag.year)
                raw["year"] = str(tag.year)
            if tag.duration:
                # Format duration into HH:MM:SS
                secs = int(tag.duration)
                duration_str = f"{secs // 3600:02d}:{(secs % 3600) // 60:02d}:{secs % 60:02d}"
                specific["Duration"] = f"{duration_str} ({secs} seconds)"
                raw["duration_seconds"] = tag.duration
            if tag.bitrate:
                specific["Bitrate"] = f"{int(tag.bitrate)} kbps"
                raw["bitrate_kbps"] = tag.bitrate
            if tag.samplerate:
                specific["Sample Rate"] = f"{int(tag.samplerate)} Hz"
                raw["samplerate_hz"] = tag.samplerate
            if tag.channels:
                specific["Channels"] = "Stereo (2)" if tag.channels == 2 else ("Mono (1)" if tag.channels == 1 else str(tag.channels))
                raw["channels"] = tag.channels
    except Exception as e:
        raw["tinytag_error"] = str(e)
        
    # 2. Try Hachoir parser to extract rich media headers (including codecs and video size)
    try:
        parser = createParser(filepath)
        if parser:
            with parser:
                metadata = extractMetadata(parser)
                if metadata:
                    for line in metadata.exportPlaintext():
                        if ":" in line:
                            k, v = line.split(":", 1)
                            key = k.strip()
                            val = v.strip()
                            raw[f"hachoir_{key.lower().replace(' ', '_')}"] = val
                            
                            # Promote key values if not already set by tinytag
                            if key == "Duration" and "Duration" not in specific:
                                specific["Duration"] = val
                            elif key == "Bitrate" and "Bitrate" not in specific:
                                specific["Bitrate"] = val
                            elif key == "Image width":
                                specific["Video Width"] = val
                            elif key == "Image height":
                                specific["Video Height"] = val
                            elif key == "Video codec":
                                specific["Video Codec"] = val
                            elif key == "Audio codec":
                                specific["Audio Codec"] = val
                            elif key == "Frame rate":
                                specific["Frame Rate"] = val
                            elif key == "Creation date" and "Creation Date" not in specific:
                                specific["Creation Date"] = val
                                
        # Construct resolution if width/height are found
        if "Video Width" in specific and "Video Height" in specific:
            specific["Resolution"] = f"{specific['Video Width']} x {specific['Video Height']}"
            
    except Exception as e:
        raw["hachoir_error"] = str(e)
        
    if not specific:
        specific["Info"] = "No standard audio/video metadata tags found, but file was analyzed."
        
    return specific, raw

def extract_metadata(filepath, original_filename):
    """
    Main forensic metadata extractor. Determines file category and calls proper parser.
    Consolidates data into general info, hashes, format-specific details, and raw dumps.
    """
    if not os.path.exists(filepath):
        return None
        
    # 1. System Metadata
    size_bytes = os.path.getsize(filepath)
    size_formatted = format_size(size_bytes)
    
    mime_type, _ = mimetypes.guess_type(original_filename)
    if not mime_type:
        mime_type = "application/octet-stream"
        
    category = get_file_category(mime_type, original_filename)
    
    # System timestamps formatted nicely
    created_time = os.path.getctime(filepath)
    modified_time = os.path.getmtime(filepath)
    
    created_str = datetime.fromtimestamp(created_time).strftime("%Y-%m-%d %H:%M:%S")
    modified_str = datetime.fromtimestamp(modified_time).strftime("%Y-%m-%d %H:%M:%S")
    
    metadata = {
        "general": {
            "filename": original_filename,
            "filepath": filepath,
            "size_bytes": size_bytes,
            "size_formatted": size_formatted,
            "mime_type": mime_type,
            "created": created_str,
            "modified": modified_str,
            "category": category,
            "extracted_at": datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        },
        "hashes": {},
        "specific": {},
        "raw": {}
    }
    
    # 2. Format Specific Analysis
    if category == 'image':
        specific, raw = extract_image_metadata(filepath)
        metadata["specific"] = specific
        metadata["raw"] = raw
    elif category == 'document' and original_filename.lower().endswith('.pdf'):
        specific, raw = extract_pdf_metadata(filepath)
        metadata["specific"] = specific
        metadata["raw"] = raw
    elif category == 'document' and original_filename.lower().endswith('.docx'):
        specific, raw = extract_docx_metadata(filepath)
        metadata["specific"] = specific
        metadata["raw"] = raw
    elif category in ['audio', 'video']:
        specific, raw = extract_audio_video_metadata(filepath)
        metadata["specific"] = specific
        metadata["raw"] = raw
    else:
        metadata["specific"] = {
            "Message": "Unsupported specific metadata parser, only general system info is available."
        }
        
    return metadata
